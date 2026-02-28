#!/usr/bin/env python3
"""
Concatenate Markdown files and export to Word document.
Reads all step Markdown files from Export/, concatenates them, and creates Erfindung-Check.docx.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

MODULE_DIR = Path(__file__).resolve().parent.parent
EXPORT_DIR = MODULE_DIR / "Export"

STEP_FILES = [
    "00-Einleitung.md",
    "01-Kategorie-waehlen.md",
    "02-Markt-scannen.md",
    "03-Prior-Art-suchen.md",
    "04-Treffer-bewerten.md",
    "05-Kern-schaerfen.md",
    "06-Ergebnis-ableiten.md",
]


def add_formatted_paragraph(doc, text):
    """Add paragraph with basic bold/italic formatting."""
    p = doc.add_paragraph()
    # Simple split by ** and *
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1])
            r.italic = True
        else:
            p.add_run(part)
    return p


def main():
    # Concatenate Markdown files
    full_md_parts = [
        "# Erfindung-Check – Arbeitsdokument\n",
        "\n*Dieses Dokument basiert auf der Webversion des Erfindung-Checks.*\n",
    ]

    for fname in STEP_FILES:
        path = EXPORT_DIR / fname
        if path.exists():
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
            full_md_parts.append(content)
            full_md_parts.append("\n\n")

    full_md = "".join(full_md_parts)

    def remove_duplicate_headings(md):
        lines = md.split("\n")
        result = []
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if (stripped.startswith("## ") or stripped.startswith("### ")) and not stripped.startswith("#### "):
                prefix_len = 4 if stripped.startswith("### ") else 3
                title = stripped[prefix_len:].strip()
                j = i + 1
                while j < len(lines) and lines[j].strip() == "":
                    j += 1
                next_line = lines[j].strip() if j < len(lines) else ""
                is_duplicate = next_line == title or (": " in title and next_line == title.split(": ", 1)[1].strip())
                if is_duplicate:
                    result.append(line)
                    i = j + 1
                    continue
            result.append(line)
            i += 1
        return "\n".join(result)

    full_md = remove_duplicate_headings(full_md)

    # Write concatenated Markdown
    full_md_path = EXPORT_DIR / "Erfindung-Check-Vollstaendig.md"
    with open(full_md_path, "w", encoding="utf-8") as f:
        f.write(full_md)
    print(f"  Wrote {full_md_path.name}")

    # Write Notion-ready version (remove internal file links, clean for paste)
    lines = full_md.split("\n")
    notion_lines = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            j = i + 1
            while j < len(lines) and lines[j].strip() == "":
                j += 1
            if j < len(lines) and re.match(r"^\*Nächster Schritt:.*\*$", lines[j].strip()):
                notion_lines.append("")  # keep one blank line as separator
                i = j + 1
                continue
        if re.match(r"^\*Nächster Schritt:.*\*$", line.strip()):
            i += 1
            continue
        notion_lines.append(line)
        i += 1
    notion_md = "\n".join(notion_lines)
    notion_path = EXPORT_DIR / "Erfindung-Check-Notion.md"
    with open(notion_path, "w", encoding="utf-8") as f:
        f.write(notion_md)
    print(f"  Wrote {notion_path.name}")

    # Convert to Word - simple line-by-line
    doc = Document()
    lines = full_md.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped.startswith("*Nächster Schritt") or stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("|") and "|" in stripped and "---" not in stripped:
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                doc.add_paragraph(" | ".join(cells))
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Pt(24)
        elif stripped.startswith("- ") or re.match(r"^\d+\. ", stripped):
            content = re.sub(r"^[-*]\s+", "", stripped) or re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(content, style="List Bullet")
        else:
            add_formatted_paragraph(doc, stripped)

        i += 1

    docx_path = EXPORT_DIR / "Erfindung-Check.docx"
    doc.save(docx_path)
    print(f"  Wrote {docx_path.name}")

    print(f"\nExported to {EXPORT_DIR}")


if __name__ == "__main__":
    main()

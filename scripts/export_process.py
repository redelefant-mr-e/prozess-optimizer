#!/usr/bin/env python3
"""
Shared export logic: HTML wizard → Markdown files, Notion, Word.
Used by all process modules (erfindung-check, patent-next-steps, erfindungsbeschreibung).
"""

import re
import subprocess
import sys
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt


def remove_duplicate_headings(md: str) -> str:
    """Remove lines that duplicate the preceding ## or ### header."""
    lines = md.split("\n")
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if (stripped.startswith("## ") or stripped.startswith("### ")) and not stripped.startswith("#### "):
            prefix_len = 4 if stripped.startswith("### ") else 3
            title = stripped[prefix_len:].strip()
            j = i + 1
            while j < len(lines) and lines[j].strip() == "":
                j += 1
            next_line = lines[j].strip() if j < len(lines) else ""
            is_duplicate = next_line == title
            if not is_duplicate and ": " in title:
                is_duplicate = next_line == title.split(": ", 1)[1].strip()
            if is_duplicate:
                result.append(line)
                i = j + 1
                continue
        result.append(line)
        i += 1
    return "\n".join(result)


def slugify(title):
    """Convert step title to filename slug."""
    replacements = {" ": "-", "ä": "ae", "ö": "oe", "ü": "ue", "ß": "ss", ":": ""}
    s = title.lower()
    for k, v in replacements.items():
        s = s.replace(k, v)
    return re.sub(r"[^a-z0-9-]", "", s)


def expand_tooltips(soup):
    for wrap in soup.find_all(class_="tooltip-wrap"):
        tooltip_span = wrap.find(class_="tooltip-text")
        if tooltip_span:
            tooltip_text = tooltip_span.get_text(strip=True)
            term = "".join(c.strip() for c in wrap.children if isinstance(c, str)).strip()
            if not term:
                term = wrap.get_text(strip=True).replace(tooltip_text, "").strip() or tooltip_text.split("=")[0].strip()
            replacement = f"{term} ({tooltip_text}) "
        else:
            replacement = wrap.get_text(strip=True)
        wrap.replace_with(replacement)
    return soup


def element_to_markdown(el):
    if el is None:
        return ""
    if isinstance(el, str):
        return el
    text_parts = []
    for child in el.children:
        if isinstance(child, str):
            text_parts.append(child)
            continue
        if child.name == "a":
            href = child.get("href", "#")
            text_parts.append(f"[{child.get_text(strip=True)}]({href})")
        elif child.name == "strong":
            text_parts.append(f"**{child.get_text(strip=True)}**")
        elif child.name == "em":
            text_parts.append(f"*{child.get_text(strip=True)}*")
        elif child.name == "code":
            text_parts.append(f"`{child.get_text(strip=True)}`")
        elif child.name == "pre":
            text_parts.append("\n```\n" + child.get_text() + "\n```\n")
        elif child.name == "br":
            text_parts.append("\n")
        elif child.name in ("h1", "h2", "h3", "h4"):
            level = int(child.name[1])
            text_parts.append("\n" + "#" * level + " " + child.get_text(strip=True) + "\n")
        elif child.name == "p":
            text_parts.append("\n" + element_to_markdown(child) + "\n")
        elif child.name == "ul":
            items = ["- " + element_to_markdown(li).strip() for li in child.find_all("li", recursive=False)]
            text_parts.append("\n" + "\n".join(items) + "\n")
        elif child.name == "ol":
            items = [f"{i}. " + element_to_markdown(li).strip() for i, li in enumerate(child.find_all("li", recursive=False), 1)]
            text_parts.append("\n" + "\n".join(items) + "\n")
        elif child.name == "blockquote":
            inner = element_to_markdown(child).strip()
            text_parts.append("\n" + "\n".join("> " + line for line in inner.split("\n")) + "\n")
        elif child.name == "table":
            text_parts.append(table_to_markdown(child))
        elif child.name == "details":
            summary = child.find("summary")
            summary_text = summary.get_text(strip=True) if summary else "Details"
            inner = element_to_markdown(child).strip()
            text_parts.append(f"\n**{summary_text}**\n{inner}\n")
        elif child.name == "div" and "wizard-table-wrap" in child.get("class", []):
            table = child.find("table")
            if table:
                text_parts.append(table_to_markdown(table))
        elif child.name == "div" and "limitations" in child.get("class", []):
            parts = ["> **⚠️ Wichtige Grenze**"]
            for p in child.find_all("p"):
                parts.append("> " + element_to_markdown(p).replace("\n", "\n> "))
            for ul in child.find_all("ul"):
                for li in ul.find_all("li", recursive=False):
                    parts.append("> - " + element_to_markdown(li).replace("\n", " "))
            text_parts.append("\n" + "\n".join(parts) + "\n")
        elif child.name in ("label", "textarea"):
            continue
        else:
            text_parts.append(element_to_markdown(child))
    return "".join(text_parts).strip()


def table_to_markdown(table):
    rows = []
    for tr in table.find_all("tr"):
        cells = [element_to_markdown(cell).replace("\n", " ").strip() for cell in tr.find_all(["th", "td"])]
        rows.append(cells)
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    for row in rows:
        while len(row) < col_count:
            row.append("")
    lines = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join("---" for _ in rows[0]) + " |"]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n" + "\n".join(lines) + "\n"


def block_to_markdown(block):
    parts = []
    label_el = block.find(class_="block-label")
    if label_el:
        label_el.decompose()
    classes = block.get("class", [])
    is_limitations = "limitations" in classes
    is_handoff = "handoff" in block.get("class", [])

    if is_limitations:
        parts.append("\n> **⚠️ Wichtige Grenze**\n")
        for p in block.find_all("p"):
            parts.append("> " + element_to_markdown(p) + "\n")
        for ul in block.find_all("ul"):
            for li in ul.find_all("li", recursive=False):
                parts.append("> - " + element_to_markdown(li) + "\n")
        return "\n".join(parts)

    for child in block.children:
        if isinstance(child, str) and not child.strip():
            continue
        if hasattr(child, "name") and child.name in ("div", "span") and "block-label" in child.get("class", []):
            continue
        content = element_to_markdown(child) if hasattr(child, "children") else str(child)
        if content.strip():
            parts.append(content)

    result = "\n".join(parts).strip()
    if "wizard-block-action" in classes:
        h3 = block.find("h3")
        section = h3.get_text(strip=True) if h3 else "Deine Aufgabe"
        result = f"## {section}\n\n" + result
        if h3:
            result = re.sub(r"^###\s*" + re.escape(h3.get_text(strip=True)) + r"\s*\n+", "", result, count=1)
    elif "wizard-block-info" in classes:
        h3 = block.find("h3")
        if h3:
            h3_text = h3.get_text(strip=True)
            result = f"## {h3_text}\n\n" + result
            result = re.sub(r"^###\s*" + re.escape(h3_text) + r"\s*\n+", "", result, count=1)
    elif "wizard-block-example" in classes:
        h3 = block.find("h3")
        if h3:
            h3_text = h3.get_text(strip=True)
            result = f"## Beispiel: {h3_text}\n\n" + result
            result = re.sub(r"^###\s*" + re.escape(h3_text) + r"\s*\n+", "", result, count=1)
        else:
            result = "## Beispiel\n\n" + result
    elif "wizard-block-tip" in classes:
        h3 = block.find("h3")
        if h3:
            h3_text = h3.get_text(strip=True)
            result = f"## Tipp: {h3_text}\n\n" + result
            result = re.sub(r"^###\s*" + re.escape(h3_text) + r"\s*\n+", "", result, count=1)
        else:
            result = "## Tipp\n\n" + result
    elif is_handoff:
        result = re.sub(r"^#+\s*Du bringst in den nächsten Schritt mit\s*\n+", "", result, count=1)
        result = "## Du bringst in den nächsten Schritt mit\n\n" + result
    elif "wizard-block-transfer" in str(classes) or "wizard-block-input" in str(classes):
        pass
    else:
        h3 = block.find("h3")
        if h3:
            h3_text = h3.get_text(strip=True)
            result = f"## {h3_text}\n\n" + result
            result = re.sub(r"^###\s*" + re.escape(h3_text) + r"\s*\n+", "", result, count=1)

    return result


def process_step(step_el, step_index, next_filename, expand_tooltips_fn):
    title_el = step_el.find(class_="wizard-step-title")
    title = title_el.get_text(strip=True) if title_el else f"Schritt {step_index}"
    if title_el:
        expand_tooltips_fn(title_el)
        title = title_el.get_text(strip=True)
        title = re.sub(r"\)([a-zA-ZäöüÄÖÜß])", r") \1", title)

    md_parts = [f"# Schritt {step_index}: {title}\n"]
    def is_block(classes):
        if not classes:
            return False
        c = classes if isinstance(classes, list) else [classes]
        return any("wizard-block" in str(x) or x in ("handoff", "limitations") for x in c)

    for block in step_el.find_all("div", class_=is_block, recursive=False):
        if "wizard-table-wrap" in block.get("class", []):
            continue
        expand_tooltips_fn(block)
        content = block_to_markdown(block)
        if content.strip():
            md_parts.append(content)
            md_parts.append("")

    if next_filename:
        md_parts.append("\n---\n")
        md_parts.append(f"*Nächster Schritt: [{next_filename}]({next_filename})*")

    return "\n".join(md_parts).strip()


def export_process(html_path: Path, export_dir: Path, step_files: list, process_title: str, output_base: str = None):
    """Export a process from HTML to Markdown, Notion, and Word."""
    output_base = output_base or slugify(process_title)
    export_dir.mkdir(parents=True, exist_ok=True)

    with open(html_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    steps = sorted(soup.find_all("div", class_="wizard-step", attrs={"data-step": True}), key=lambda s: int(s.get("data-step", 0)))

    for i, step_el in enumerate(steps):
        next_file = step_files[i + 1] if i + 1 < len(step_files) else None
        md_content = process_step(step_el, i, next_file, expand_tooltips)

        out_path = export_dir / step_files[i]
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"  Wrote {out_path.name}")

    # Concatenate
    full_md_parts = [f"# {process_title} – Arbeitsdokument\n", f"\n*Dieses Dokument basiert auf der Webversion.*\n"]
    for fname in step_files:
        path = export_dir / fname
        if path.exists():
            full_md_parts.append(path.read_text(encoding="utf-8"))
            full_md_parts.append("\n\n")
    full_md = remove_duplicate_headings("".join(full_md_parts))

    # Vollstaendig
    (export_dir / f"{output_base}-Vollstaendig.md").write_text(full_md, encoding="utf-8")
    print(f"  Wrote {output_base}-Vollstaendig.md")

    # Notion (remove internal links)
    lines = full_md.split("\n")  # full_md already has duplicates removed
    notion_lines = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            j = i + 1
            while j < len(lines) and lines[j].strip() == "":
                j += 1
            if j < len(lines) and re.match(r"^\*Nächster Schritt:.*\*$", lines[j].strip()):
                notion_lines.append("")
                i = j + 1
                continue
        if re.match(r"^\*Nächster Schritt:.*\*$", line.strip()):
            i += 1
            continue
        notion_lines.append(line)
        i += 1
    (export_dir / f"{output_base}-Notion.md").write_text("\n".join(notion_lines), encoding="utf-8")
    print(f"  Wrote {output_base}-Notion.md")

    # Word
    doc = Document()
    for line in full_md.split("\n"):
        stripped = line.strip()
        if not stripped or stripped.startswith("*Nächster Schritt") or stripped == "---":
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("|") and "|" in stripped and "---" not in stripped:
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                doc.add_paragraph(" | ".join(cells))
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Pt(24)
        elif stripped.startswith("- ") or re.match(r"^\d+\. ", stripped):
            content = re.sub(r"^[-*]\s+", "", stripped) or re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(content, style="List Bullet")
        else:
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", stripped):
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    r = p.add_run(part[1:-1])
                    r.italic = True
                else:
                    p.add_run(part)

    doc.save(str(export_dir / f"{output_base}.docx"))
    print(f"  Wrote {output_base}.docx")

    print(f"\nExported to {export_dir}")


